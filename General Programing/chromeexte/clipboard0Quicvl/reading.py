text = """
Pytts3  text to speech conversion library.
Offline
Supports TTS engines including Sapi5, nsss, and espeak

"""
import os
from os.path import join
import random
import math
import time

from docx import Document
from docx.shared import Inches

import re
import pyttsx3

import pyperclip as pc

import traceback

def notes(text):
    PATH = "C:/Users/user/Documents/Books Reading/Notes/"
    document = Document()
    paragraph = document.add_paragraph(text[:50])
    paragraph.bold = True
    paragraph.add_run(text[51:])
    text = text[:70].splitlines()[0]
    text = text.replace("?", "")
    PATH = PATH + "{}.docx".format(text)
    try:
        document.save(PATH)
    except:
        PATH = "1 {}".format(int(time.time())) + PATH
        document.save(PATH)

text2 = pc.paste().rstrip("\n")
engine = pyttsx3.init()
voices = engine.getProperty('voices')
rate = engine.getProperty('rate')
engine.setProperty('rate', rate)
engine.setProperty('voice', voices[0].id)
#  Clean ans noramalization

replacer1 = ['[1]', '’','‘', '•', '“', '”', "Go to:"]
replacer3 = ['’re', '']
replacer2 = ['–']
text = text2
for i in replacer1:
    text = text.replace(i, '')
for i in replacer2:
    text = text.replace(i,
     ' ')
text = text.replace("#", ' ')
text = text.replace("(s)", 's')
text = text.replace("_", ' ')
text = text.replace("[…]", '.')
text = text.replace("—", ' ')
text = text.replace("XMLHttpRequest", 'XHR ')
# text = ' '.join(text2.split())
text = re.sub(r'\[\d\]', '', text)
# text = replace('\{\\displaystyle \\gamma .\}\\gamma .', ' gamma ')
text = re.sub('’s', 's', text)
notes(text)
text = text.lower()
text = text.replace('we’ll', 'we will')
text = text.replace('you’ll', 'you will')
text = text.replace('• ', '( )')


#Meta_Data
# Say
start = time.time()
engine.say(text)
engine.runAndWait()
stop = time.time()
txt_length = len(text)
with open("metrics.txt", "a") as file:
    read_date = " , ".join(time.ctime().split()[1:3])
    read_year = time.localtime().tm_year
    read_hour = time.localtime().tm_hour
    read_time = stop - start
    tupled = '{}, {}, {}, {}, {} \n'.format(read_year, read_date, read_hour, int(read_time), txt_length)
    print(tupled)
    file.write(tupled)
    file.close()

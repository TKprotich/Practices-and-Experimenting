""" Documentation Hub"""
import time
import datetime as dt

import os
from os.path import join
import random
import math

from docx import Document
from docx.shared import Inches

from pydub import AudioSegment
from pydub.playback import play

from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager

import traceback

sound0 = AudioSegment.from_mp3(r'C:\Users\user\Documents\notification\a.mp3')

sound1 = AudioSegment.from_mp3(r'C:\Users\user\Documents\notification\notification5.mp3')
sound2 = AudioSegment.from_mp3(r'C:\Users\user\Documents\Songs\Jesus-Lover-of-My-soul.mp3')

########
PATH = r'C:\Users\user\Documents\notification\workout'
walert1 = [i for i in os.listdir(PATH+"\walert1") if i.endswith(".m4a")]
walert2 = [i for i in os.listdir(PATH+"\walert2") if i.endswith(".m4a")]
walert3 = [i for i in os.listdir(PATH+"\walert3") if i.endswith(".m4a")]
walert4 = [i for i in os.listdir(PATH+"\walert4") if i.endswith(".m4a")]
walert5 = [i for i in os.listdir(PATH+"\walert5") if i.endswith(".m4a")]
walert6 = [i for i in os.listdir(PATH+"\walert6") if i.endswith(".m4a")]
walert7 = [i for i in os.listdir(PATH+"\walert7") if i.endswith(".m4a")]
walertlist = [walert1, walert2, walert3, walert4, walert5, walert6, walert7]
###
def excercise():
    l =0
    for walert in walertlist:
        l +=1
        for alert in walert:
            pathtovoice = PATH + "\walert{}".format(l) + "\\" + alert
            selected_voice = AudioSegment.from_file(pathtovoice, format='m4a')
            play(sound0)
            time.sleep(1)
            play(selected_voice)
            pathtosong =PATH + "\song{}.mp3".format(l//3 +1)
            song = AudioSegment.from_mp3(pathtosong)
            play(song)
            # Pause other things and do this????????????????????????????????????

def randomquestion():
    pass
    
def notes():
    PATH = "C:/Users/user/Documents"
    if input('New or Old?: ').lower() == "new":
        document = Document()
        name = input('What is the Name of the Notebook: ')
    else:
        name = input('Notes Name: ').capitalize()
        path = PATH +"/"+ name + ".docx"
        document = Document(path)
    topic = input('What is the Topic: ')
    document.add_heading(topic, 0)
    document.add_heading(topic + ' Notes')
    note = input('Type a note:\n')
    document.add_paragraph(note, style='List Bullet')
    while True:
        try:
            note = input('Type a another:\n')
            document.add_paragraph(note, style='List Bullet')
        except:
            document.save(name.capitalize() + '.docx')
            break
            main()
def entertainment():
    PATH = r"C:\Users\user\Music"
    songs = []
    for root, dirs, files in os.walk(PATH+"\\all"):
        songs.append([join(root, name) for name in files if name.endswith(".mp3")])
    while True:
        try:
            pathtosong = random.choice(songs[0])
            selected_song = AudioSegment.from_mp3(pathtosong)
            sliced = selected_song[7000:] # get a slice from 5 to 10 seconds of an mp3
            play(selected_song)
        except:
            main()
def openbrowser():
    global driver
    driver = webdriver.Chrome(ChromeDriverManager().install())
    driver.get("https://web.whatsapp.com")
    print("Scan QR Code, And then Enter")
    input()
    print("Logged In")

def closebrowser():
    driver.quit()
def sendmessage(msg, contact):
    try:
        user = driver.find_element_by_xpath("//span[@title='{}']".format(contact))
        user.click()
        time.sleep(2)

        selected_contact = driver.find_element_by_xpath("//span[@title='"+contact+"']")
        selected_contact.click()

        inp_xpath = '//div[@class="_3FRCZ copyable-text selectable-text"][@contenteditable="true"][@data-tab="1"]'
        input_box = driver.find_element_by_xpath(inp_xpath)
        time.sleep(2)
        input_box.send_keys(msg + Keys.ENTER)
        time.sleep(2)
        return("Success")
    except Exception:
        traceback.print_exc()
        return("error")
def studytrack():
    print("studying.................")
    i = 0
    while True:
        i = i + 1
        time.sleep(5*60)
        if i ==1:
            print(i*5, " Minute Spend")
        else:
            print(i*5 , " Minutes Spend")
        try:
            if i*5 % 120 == 0:
                play(sound1)
                #play(sound2)
                book = input('Enter the book that you read: ')
                description = input('Description: What chapter? What page? ')
                feeling = input('Reflection: How do I feel about this book? ')
                datetime_object = dt.datetime.now()
                value = input('Give a rating on the value of this reading High, Medium, Low, None):  ')
                readinglogs = open(r'C:\Users\user\Documents\Books Reading\reading.txt', 'a+')
                readinglogs.write('{}, {}, {}, {} \n'.format(book, description, datetime_object, value))
                readinglogs.close()
                #Estimate time when the book will be finished
                pagesunread = int(input('How many pages of the book did you read Remains? '))
                pagesread = float(input('How many pages of the book did you read? '))
                rate_of_reading = 25/pagesread
                remaining_time = math.ceil((rate_of_reading*pagesunread)/360)
                smessage = '25 minutes done! I read {}:\n   -Description: {}.\n  -Date: {}.\n    -Value: {}.\n   -It might take you {} days to finish for 6 hour reading a day'.format(book, description, datetime_object, value, remaining_time )
                print(smessage)
                notes()
            else:
                play(sound1)
        except Exception:
            traceback.print_exc()
def actionlog():
    #60*60*2
    i = 0
    start = time.time()
    stop = start + (3)
    while time.time() <= stop:
        try:
            entertainment()
            play(sound1)
            play(sound2)
            activity = input('Activity: ')
            description = input('Description')
            feeling = input('How do I feel about this work? ')
            datetime_object = dt.datetime.now()
            value = input('Give a rating on the value of this work High, Medium, Low, None):  ')
            activitylogs = open(r'C:\Users\user\Documents\Books Reading\activitylogs.txt', 'a+')
            activitylogs.write('{}, {}, {}, {} \n'.format(activity, description, datetime_object, value))
            activitylogs.close()
            sendmessage('25 minutes done! I have done {} \n, Description: {} \n, Date: {} \n, Value: {} \n'.format(activity, description, datetime_object, value), "My Messages")
            if input('Do you intend to keep working? ').lower()== 'yes':
                continue
            else:
                break
                main()
        except:
            main()
def main():
    try:
        openbrowser()
        sendmessage("Good to go!", "My Messages")
    except:
        print("Internet Problem")
    while True:
        try:
            if input('Do you intend to read? ').lower() == 'yes':
                studytrack()
            elif input('Do you intend to Work? ').lower() == 'yes':
                actionlog()
            elif input('Do you intend to Write a note? ').lower() == 'yes':
                notes()
            elif input('Do you intend to do excercise? ').lower() == 'yes':
                excercise()
            else:
                entertainment()
        except Exception:
            traceback.print_exc()
            main()
if __name__ == "__main__":
 main()



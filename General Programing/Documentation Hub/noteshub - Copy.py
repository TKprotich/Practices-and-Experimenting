import os
from os.path import join
import random
import math

from docx import Document
from docx.shared import Inches


import traceback
class notes:
    def __init__(self, topic):
        self.topic = topic
        self.title = topic + ' Notes'
    def registernotebook(self, noteslist):
        document.add_heading(self.topic, 0)
        document.add_heading(self.title)
        for note in self.noteslist:
            document.add_paragraph(note, style='List Bullet')
    def notesdialogue(self, docdir):
        PATH = "C:/Users/user/Documents"
        if input('New or Old?: ').lower() == "new":
            document = Document()
            name = input('What is the Name of the Notebook: ')
        else:
            name = input('Notes Name: ').capitalize()
            path = PATH +"/"+ name + ".docx"
            document = Document(path)
        topic = input('What is the Topic: ')
        document.add_heading(topic, 0)
        document.add_heading(topic + ' Notes')
        note = input('Type a note:\n')
        document.add_paragraph(note, style='List Bullet')
        while True:
            try:
                note = input('Type a another:\n')
                document.add_paragraph(note, style='List Bullet')
            except:
                document.save(r"C:\Users\user\Documents\Books Reading\Notes\\" + name.capitalize() + '.docx')
                break
if __name__ == "__main__":
    notes()

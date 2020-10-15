import os
from os.path import join
import random
import math

from docx import Document
from docx.shared import Inches


import traceback

def notes():
    PATH = r"C:\Users\user\Documents\Books Reading\Notes\General Notes.docx"
    topic = input('What is the Topic: ')
    document.add_heading(topic, 0)
    document.add_heading(topic + ' Notes')
    note = input('Type a note:\n')
    document.add_paragraph(note, style='List Bullet')
    while True:
        try:
            note = input('Type a another:\n')
            document.add_paragraph(note, style='List Bullet')
        except:
            document.save(r"C:\Users\user\Documents\Books Reading\Notes\\" + name.capitalize() + '.docx')
            break
if __name__ == "__main__":
    notes()

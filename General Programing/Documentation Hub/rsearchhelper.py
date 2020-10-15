import os
from os.path import join
import random
import math

from docx import Document
from docx.shared import Inches


import traceback

def notes():
    PATH = "C:\\Users\\user\\Documents\\2020-Kalkaal Speciality Hospital\\Research and Development\\Covid19.docx"
    document = Document(PATH)
    topic = input('What is the Topic: ')
    text2 = pc.paste().rstrip("\n")
    #  Clean ans noramalization
    replacer1 = ['’','‘', '•', '“', '”']
    replacer2 = ['–', "#", "[…]", "→", "_", "—", "", "‐"]
    # text = ' '.join(text2.split())
    for i in replacer1:
        text = text.replace(i, '')
    for i in replacer2:
        text = text.replace(i, ' ')
    text = text.replace("XMLHttpRequest", 'XHR ')
    text = ' '.join(text.splitlines())

    #Writing
    if topic !="":
        document.add_heading(topic)
    document.add_paragraph(text)
         
    document.save(PATH)
            
if __name__ == "__main__":
    notes()
